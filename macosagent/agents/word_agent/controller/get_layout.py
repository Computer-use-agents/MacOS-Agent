from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.shape import CT_Picture
from docx.oxml.table import CT_Tbl
import zipfile
import base64
from collections import OrderedDict

def safe_round(value, decimals=1):
    """安全地进行四舍五入，处理None值"""
    return round(value, decimals) if value is not None else None

def parse_word_document(file_path):
    """解析Word文档，返回按顺序排列的所有元素"""
    doc = Document(file_path)
    elements = []
    
    # 提取文档中的图片资源
    image_resources = extract_images(file_path)
    image_counter = 1
    
    # 遍历文档主体中的每个XML元素
    for element in doc.element.body:
        element_data = OrderedDict()
        
        # 处理段落
        if element.tag.endswith('p'):
            try:
                para = doc.paragraphs[len([e for e in elements if e.get('type') == 'paragraph'])]
                element_data['type'] = 'paragraph'
                element_data['text'] = para.text
                element_data['style'] = para.style.name if para.style else None
                element_data['alignment'] = str(para.alignment).split('.')[-1].lower() if para.alignment else None
                
                # 段落格式
                pf = para.paragraph_format
                element_data['format'] = OrderedDict([
                    ('left_indent', safe_round(pf.left_indent.inches, 3) if pf.left_indent else None),
                    ('right_indent', safe_round(pf.right_indent.inches, 3) if pf.right_indent else None),
                    ('first_line_indent', safe_round(pf.first_line_indent.inches, 3) if pf.first_line_indent else None),
                    ('space_before', safe_round(pf.space_before.pt, 1) if pf.space_before else None),
                    ('space_after', safe_round(pf.space_after.pt, 1) if pf.space_after else None),
                    ('line_spacing', safe_round(pf.line_spacing, 1) if pf.line_spacing is not None else None),
                    ('line_spacing_rule', str(pf.line_spacing_rule).split('.')[-1].lower() if pf.line_spacing_rule else None)
                ])
                
                # 处理段落中的runs
                element_data['runs'] = []
                for run in para.runs:
                    run_data = OrderedDict()
                    run_data['text'] = run.text
                    run_data['bold'] = run.bold
                    run_data['italic'] = run.italic
                    run_data['underline'] = run.underline
                    run_data['font'] = OrderedDict([
                        ('name', run.font.name),
                        ('size', safe_round(run.font.size.pt, 1) if run.font and run.font.size else None),
                        ('color', rgb_to_hex(run.font.color.rgb) if run.font and run.font.color and run.font.color.rgb else None),
                        ('highlight', rgb_to_hex(run.font.highlight_color) if run.font and run.font.highlight_color else None)
                    ])
                    element_data['runs'].append(run_data)
                
                # 检查段落中是否包含内联图片
                for run in para.runs:
                    for inline in run.element.iter():
                        if inline.tag.endswith('drawing'):
                            for pic in inline.iter():
                                if pic.tag.endswith('blip'):
                                    rId = pic.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                    if rId and rId in image_resources:
                                        element_data['inline_image'] = OrderedDict([
                                            ('image_id', f"img_{image_counter}"),
                                            ('data', image_resources[rId]),
                                            ('format', 'base64'),
                                            ('description', '')
                                        ])
                                        image_counter += 1
            except Exception as e:
                print(f"error when processing: {str(e)}")
                continue
        
        # 处理表格
        elif element.tag.endswith('tbl'):
            try:
                table = doc.tables[len([e for e in elements if e.get('type') == 'table'])]
                element_data['type'] = 'table'
                element_data['rows'] = len(table.rows)
                element_data['columns'] = len(table.columns)
                element_data['alignment'] = str(table.alignment).split('.')[-1].lower() if table.alignment else None
                element_data['style'] = table.style.name if table.style else None
                
                # 表格列宽
                element_data['column_widths'] = [
                    safe_round(col.width.inches, 3) if col.width else None 
                    for col in table.columns
                ]
                
                # 处理单元格
                element_data['cells'] = []
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        cell_data = OrderedDict([
                            ('row', i),
                            ('col', j),
                            ('text', cell.text),
                            ('width', safe_round(cell.width.inches, 3) if cell.width else None),
                            ('vertical_alignment', str(cell.vertical_alignment).split('.')[-1].lower() if cell.vertical_alignment else None),
                            ('paragraphs', [])
                        ])
                        
                        # 处理单元格中的段落
                        for para in cell.paragraphs:
                            para_data = OrderedDict([
                                ('text', para.text),
                                ('alignment', str(para.alignment).split('.')[-1].lower() if para.alignment else None),
                                ('runs', [])
                            ])
                            
                            # 处理单元格段落中的runs
                            for run in para.runs:
                                run_data = OrderedDict([
                                    ('text', run.text),
                                    ('bold', run.bold),
                                    ('italic', run.italic),
                                    ('underline', run.underline),
                                    ('font', OrderedDict([
                                        ('name', run.font.name if run.font else None),
                                        ('size', safe_round(run.font.size.pt, 1) if run.font and run.font.size else None),
                                        ('color', rgb_to_hex(run.font.color.rgb) if run.font and run.font.color and run.font.color.rgb else None)
                                    ]))
                                ])
                                para_data['runs'].append(run_data)
                            
                            cell_data['paragraphs'].append(para_data)
                        
                        element_data['cells'].append(cell_data)
            except Exception as e:
                print(f"处理表格时出错: {str(e)}")
                continue
        
        # 处理浮动图片
        elif hasattr(element, 'graphic'):
            element_data['type'] = 'floating_image'
            element_data['description'] = 'Floating image (needs additional processing)'
        
        if element_data:
            elements.append(element_data)
    
    return elements

def extract_images(file_path):
    """提取Word文档中的所有图片"""
    images = {}
    try:
        with zipfile.ZipFile(file_path) as z:
            for rel in z.namelist():
                if rel.startswith('word/media/'):
                    with z.open(rel) as img_file:
                        images[rel.split('/')[-1]] = base64.b64encode(img_file.read()).decode('utf-8')
    except Exception as e:
        print(f"提取图片时出错: {str(e)}")
    return images

def rgb_to_hex(rgb_color):
    """将RGBColor对象转换为十六进制颜色代码"""
    if not rgb_color:
        return None
    return '#{:02x}{:02x}{:02x}'.format(
        rgb_color.r,
        rgb_color.g,
        rgb_color.b
    )
# 使用示例
if __name__ == "__main__":
    file_path = "input.docx"  # 替换为你的Word文档路径
    document_elements = parse_word_document(file_path)
    
    # 打印前几个元素的信息
    for i, element in enumerate(document_elements[:3]):
        print(f"\n元素 {i+1}: {element['type']}")
        print("-"*40)
        for key, value in element.items():
            if key in ['runs', 'cells']:
                print(f"{key}: [{len(value)} items]")
            elif isinstance(value, dict):
                print(f"{key}:")
                for sub_key, sub_value in value.items():
                    print(f"  {sub_key}: {sub_value}")
            else:
                print(f"{key}: {value}")
    
    # document_elements 就是包含所有元素的列表