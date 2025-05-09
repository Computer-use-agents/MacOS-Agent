import os
from pathlib import Path

from docx import Document
from docx2pdf import convert
from pdf2docx import Converter


def convert_file_function(input_path: str, output_path: str) -> None:
    """
    智能文档转换函数
    支持转换类型：
    - docx → pdf
    - pdf → docx
    - docx ↔ txt
    
    参数：
    input_path: 输入文件路径
    output_path: 输出文件路径（自动识别格式）
    """
    # 规范化路径
    input_path = Path(input_path)
    output_path = Path(output_path)
    
    # 验证输入文件
    if not input_path.exists():
        raise FileNotFoundError(f"输入文件不存在: {input_path}")
    
    # 获取格式信息
    input_ext = input_path.suffix.lower()
    output_ext = output_path.suffix.lower()
    
    # 创建输出目录
    output_path.parent.mkdir(parents=True, exist_ok=True)
    # print(output_path)
    # 转换路由字典
    conversion_map = {
        ('.docx', '.pdf'): _convert_docx_to_pdf,
        ('.docx', '.txt'): _convert_docx_to_txt,
        ('.pdf', '.docx'): _convert_pdf_to_docx,
        ('.txt', '.docx'): _convert_txt_to_docx,
    }
    
    # 查找转换方法
    converter = conversion_map.get((input_ext, output_ext))
    if not converter:
        supported = [f"{k[0]}→{k[1]}" for k in conversion_map.keys()]
        raise ValueError(f"不支持的转换类型，当前支持：{', '.join(supported)}")
    
    # 执行转换
    converter(input_path, output_path)

# 以下是具体转换实现
def _convert_docx_to_pdf(input_path: Path, output_path: Path) -> None:
    """DOCX转PDF"""
    try:
        convert(str(input_path), str(output_path))
    except Exception as e:
        # 备用方案：使用LibreOffice命令行
        os.system(f'soffice --convert-to pdf "{input_path}" --outdir "{output_path.parent}"')

def _convert_pdf_to_docx(input_path: Path, output_path: Path) -> None:
    """PDF转DOCX"""
    cv = Converter(str(input_path))
    cv.convert(str(output_path), start=0, end=None)
    cv.close()

def _convert_docx_to_txt(input_path: Path, output_path: Path) -> None:
    """DOCX转TXT"""
    doc = Document(input_path)
    with open(output_path, 'w', encoding='utf-8') as f:
        for para in doc.paragraphs:
            f.write(para.text + '\n')

def _convert_txt_to_docx(input_path: Path, output_path: Path) -> None:
    """TXT转DOCX"""
    doc = Document()
    with open(input_path, 'r', encoding='utf-8') as f:
        doc.add_paragraph(f.read())
    doc.save(output_path)



from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.oxml.ns import qn
from typing import List, Tuple
import re
from docx.oxml import OxmlElement
from docx.oxml.shared import qn

def modify_ranges_styles_function(doc_path: str, save_path: str, 
                        paragraph_to_modify: int ,
                        start_index: int,
                        end_index: int,
                        font_name: str = None, 
                        font_size: float = None,
                        font_color: str | Tuple[int, int, int] = None,
                        left_indent_cm=None, first_line_indent_cm=None):
    
    """
    基于字符位置的精确样式修改
    
    参数:
        doc_path: 文档位置
        save_path: 存储位置
        ranges: 需要修改的区间列表，每个元素为元组：
                (段落索引从0开始, 起始位置包含, 结束位置不包含)
                示例：[(0,5,10), (1,3,7)] 表示修改：
                    - 第1段5-9字符
                    - 第2段3-6字符
        font_name: 字体
        font_size: 字大小
        font_color: 字颜色
    示例：
    modify_ranges_styles(
        "input.docx",
        "output.docx",
        ranges=[(0, 5, 10), (1, 3, 7)],  # 修改第1段5-9字符，第2段3-6字符
        font_name="微软雅黑",
        font_size=12,
        font_color="#FF0000"
    )
    """
    # ranges = eval(ranges) 
    # [{"paragraph":1,"start_number":0,"end_number":10}]]
    if paragraph_to_modify>0:
        paragraph_to_modify = paragraph_to_modify -1
    if start_index>0:
        start_index = start_index -1
    # if end_index>0:
    #     end_index = end_index -1
    ranges = [paragraph_to_modify,start_index,end_index]
    try:
        doc = Document(doc_path)
    except:
        raise IOError("The word file you want to operate is not currently available to manipulate. you need to first save and close the file. ")
    # 颜色转换（同之前实现)
    color_rgb = None
    if font_color:
        if isinstance(font_color, str):
            hex_color = font_color.strip().lstrip('#')
            if len(hex_color) != 6:
                raise ValueError("颜色格式错误，应为#RRGGBB格式")
            r, g, b = (int(hex_color[i:i+2], 16) for i in (0, 2, 4))
            color_rgb = RGBColor(r, g, b)
        elif isinstance(font_color, (tuple, list)) and len(font_color) == 3:
            if not all(0 <= c <= 255 for c in font_color):
                raise ValueError("RGB数值应在0-255之间")
            color_rgb = RGBColor(*font_color)
        else:
            raise TypeError("颜色应为#RRGGBB字符串或RGB元组")
    
    para_format = {}
    if left_indent_cm is not None:
        para_format['left_indent'] = Inches(left_indent_cm / 2.54)
    if first_line_indent_cm is not None:
        para_format['first_line_indent'] = Inches(first_line_indent_cm / 2.54)
    
    # 按段落索引分组处理区间
    # 按段落分组处理区间
    from collections import defaultdict
    para_ranges = defaultdict(list)
    ranges = [ranges]
    for para_idx, start, end in ranges:
        para_ranges[para_idx].append((start, end))

    # 处理每个段落
    for para_idx, para in enumerate(doc.paragraphs):
        if para_idx not in para_ranges:
            continue

        # 建立字符位置映射表（修正版）
        char_mapping = []
        global_pos_counter = 0
        for run in para.runs:
            for run_pos, _ in enumerate(run.text):
                char_mapping.append( (global_pos_counter, (run, run_pos)) )
                global_pos_counter += 1

        # 处理每个目标区间
        targets = set()
        for start, end in para_ranges[para_idx]:
            targets.update(range(start, end))

        # 重构段落内容（修复变量作用域问题）
        new_runs = []
        current_run = None
        for global_pos, (run, run_pos) in char_mapping:
            # 始终先获取当前字符
            char = run.text[run_pos]  # 这行是修复关键
            
            if global_pos in targets:
                # 需要修改的字符处理
                if current_run is None or not current_run['modified']:
                    current_run = {
                        'text': [char],
                        'modified': True,
                        'original_style': run
                    }
                    new_runs.append(current_run)
                else:
                    current_run['text'].append(char)
            else:
                # 保留原样的字符处理
                if current_run is None or current_run['modified']:
                    current_run = {
                        'text': [char],
                        'modified': False,
                        'original_style': run
                    }
                    new_runs.append(current_run)
                else:
                    current_run['text'].append(char)

        # 清空并重建段落（保持其他逻辑不变）
        para.clear()
        for run_info in new_runs:
            new_run = para.add_run(''.join(run_info['text']))
            if run_info['modified']:
                # 应用新样式
                if font_name:
                    new_run.font.name = font_name
                    new_run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                if font_size:
                    new_run.font.size = Pt(font_size)
                if color_rgb:
                    new_run.font.color.rgb = color_rgb
            else:
                # 复制原样式
                _copy_run_style(run_info['original_style'], new_run)

    doc.save(save_path)

def _copy_run_style(source_run, target_run):
    """安全复制Run样式（增强兼容性版）"""
    source_rPr = source_run._element.rPr
    if source_rPr is None:
        return

    target_rPr = target_run._element.get_or_add_rPr()

    # 使用更安全的元素访问方式
    # 1. 处理字体
    if source_rPr.rFonts is not None:
        target_rFonts = OxmlElement('w:rFonts')
        for attr in ['w:ascii', 'w:eastAsia', 'w:hAnsi', 'w:cs']:
            val = source_rPr.rFonts.get(qn(attr))
            if val:
                target_rFonts.set(qn(attr), val)
        target_rPr.append(target_rFonts)

    # 2. 处理字号（使用find查找元素）
    sz_element = source_rPr.find(qn('w:sz'))
    if sz_element is not None:
        target_sz = OxmlElement('w:sz')
        target_sz.set(qn('w:val'), sz_element.get(qn('w:val')))
        target_rPr.append(target_sz)

    # 3. 处理颜色（使用find查找元素）
    color_element = source_rPr.find(qn('w:color'))
    if color_element is not None:
        target_color = OxmlElement('w:color')
        for attr in ['w:val', 'w:themeColor']:
            val = color_element.get(qn(attr))
            if val:
                target_color.set(qn(attr), val)
        target_rPr.append(target_color)

    # 4. 处理基础属性（使用安全访问方式）
    target_run.bold = source_run.bold if source_run.bold is not None else False
    target_run.italic = source_run.italic if source_run.italic is not None else False



# from docx import Document
# from pptx import Presentation
# from pptx.util import Inches, Pt
# from pptx.dml.color import RGBColor
# import io
# import os
# import re
# from docx.oxml import OxmlElement
# from copy import deepcopy


# class SlideGenerator:
#     def __init__(self):
#         self.prs = Presentation()
#         self.current_slide = None
#         self.current_layout = None
#         self.content_top = Inches(1.5)
#         self.content_left = Inches(0.5)
#         self.content_width = Inches(9)
#         self.max_content_height = Inches(6.5)
#         self.current_y = self.content_top
#         self.page_count = 0

#     def create_slide(self, slide_type, title=None, level=1):
#         """创建新幻灯片并重置排版参数"""
#         layout_map = {
#             'title': 0,
#             'content': 1,
#             'section': 2
#         }
        
#         # 选择布局类型
#         if slide_type == 'title':
#             layout_idx = layout_map['title']
#         # elif level <= 2:
#         #     layout_idx = layout_map['section']
#         else:
#             layout_idx = layout_map['content']


#         try:
#             slide_layout = self.prs.slide_layouts[layout_idx]
#         except IndexError:
#             slide_layout = self.prs.slide_layouts[1]
#         # self.prs.slide_layouts[1]
#         self.current_slide = self.prs.slides.add_slide(slide_layout)
#         self.page_count += 1
#         self.current_y = self.content_top
        
#         # 设置标题样式
#         if title:
#             title_box = self.current_slide.shapes.title
#             title_box.text = title
#             self.apply_title_style(title_box, level)
            
#         return self.current_slide

#     def apply_title_style(self, title_box, level):
#         """应用标题级别样式"""
#         color_gradient = [
#             RGBColor(0, 32, 96),    # 1级标题
#             RGBColor(13, 54, 133),  # 2级标题
#             RGBColor(79, 129, 189)  # 3级+标题
#         ]
#         font_size = Pt(32 - (level * 4))
        
#         title_box.text_frame.paragraphs[0].font.bold = True
#         title_box.text_frame.paragraphs[0].font.color.rgb = color_gradient[min(level-1, 2)]
#         title_box.text_frame.paragraphs[0].font.size = font_size

#     def add_content(self, element):
#         """智能添加内容元素"""
#         required_height = self.calculate_element_height(element)
        
#         # 分页判断逻辑
#         if (self.current_y + required_height) > self.max_content_height:
#             self.create_continuation_slide()
        
#         # 实际添加内容
#         added_height = self.place_element(element)
#         self.current_y += added_height + Inches(0.2)

#     def calculate_element_height(self, element):
#         """计算元素所需高度"""
#         if element['type'] == 'text':
#             return self.calculate_text_height(element['content'])
#         elif element['type'] == 'image':
#             return element['height']
#         elif element['type'] == 'table':
#             return element['rows'] * Inches(0.4)
#         return Inches(0)

#     def calculate_text_height(self, text):
#         """精确计算文本高度"""
#         avg_char_width = Inches(0.15)  # 每个字符宽度
#         chars_per_line = int(self.content_width / avg_char_width)
#         lines = sum([len(line)//chars_per_line + 1 for line in text.split('\n')])
#         return lines * Inches(0.22)     # 每行高度

#     def create_continuation_slide(self):
#         """创建续页"""
#         current_title = self.current_slide.shapes.title.text
#         new_title = f"{current_title}"
#         self.create_slide('content', new_title, level=3)

#     def place_element(self, element):
#         """实际排版元素"""
#         if element['type'] == 'text':
#             textbox = self.current_slide.shapes.add_textbox(
#                 self.content_left, self.current_y,
#                 self.content_width, self.calculate_text_height(element['content'])
#             )
#             text_frame = textbox.text_frame
#             text_frame.word_wrap = True
#             text_frame.text = element['content']
#             self.apply_body_style(text_frame)
#             return textbox.height
        
#         elif element['type'] == 'image':
#             img = self.current_slide.shapes.add_picture(
#                 element['stream'], 
#                 self.content_left, self.current_y,
#                 width=self.content_width, height=element['height']
#             )
#             return img.height
        
#         elif element['type'] == 'table':
#             table = self.current_slide.shapes.add_table(
#                 element['rows'], element['cols'],
#                 self.content_left, self.current_y,
#                 self.content_width, element['rows'] * Inches(0.4)
#             ).table
#             self.fill_table_data(table, element['data'])
#             return table.height

#     def apply_body_style(self, text_frame):
#         """应用正文样式"""
#         for paragraph in text_frame.paragraphs:
#             paragraph.font.size = Pt(11)
#             paragraph.line_spacing = 1.2

#     def fill_table_data(self, table, data):
#         """填充表格数据"""
#         for row_idx, row_data in enumerate(data):
#             for col_idx, text in enumerate(row_data):
#                 cell = table.cell(row_idx, col_idx)
#                 cell.text = text
#                 cell.fill.solid()
#                 cell.fill.fore_color.rgb = RGBColor(242, 242, 242)

# def convert_docx_to_ppt_function(docx_path, pptx_path):
#     max_paragraph_len = 500
#     generator = SlideGenerator()
#     doc = Document(docx_path)
    
#     # 创建封面页
#     # generator.create_slide('title', " ")
    
#     current_section = None
#     content_buffer = []

#     def flush_buffer():
#         """提交缓冲内容"""
#         if content_buffer:
#             for element in content_buffer:
#                 generator.add_content(element)
#             content_buffer.clear()
#     paragraph_list = []
#     # print(1)
#     for p in doc.paragraphs:
#         # print(p)
#         start_pos = 0
#         if len(p.text) >= max_paragraph_len:
#             while len(p.text) - start_pos >= max_paragraph_len:
#                 t = deepcopy(p)
#                 t.text = p.text[start_pos:start_pos+max_paragraph_len]
#                 paragraph_list.append(t)
#                 # parse_paragraph.append(p)
#                 start_pos += max_paragraph_len
#             t = deepcopy(p)
#             t.text = p.text[start_pos:]
#             paragraph_list.append(t)
#         else:
#             paragraph_list.append(p)

#     for para in paragraph_list:
#         if is_heading(para):
#             # 提交之前的内容
#             flush_buffer()
            
#             # 处理新标题
#             level = get_heading_level(para)
#             # generator.create_slide('section' if level <=2 else 'content', para.text, level)
#             generator.create_slide('content', para.text, level)
#             current_section = para.text
#             continue
        
#         # 处理段落内容
#         element = parse_paragraph(para, doc)
#         if element:
#             content_buffer.append(element)
    
#     # 处理表格
#     for table in doc.tables:
#         content_buffer.append(parse_table(table))
    
#     # 提交最后的内容
#     flush_buffer()
    
    
#     # 保存最终文件
#     generator.prs.save(pptx_path)
#     print(f"成功生成PPT：{os.path.abspath(pptx_path)}")
    
#     # 再删除第一页！

# # 辅助工具函数
# def is_heading(paragraph):
#     style_name = paragraph.style.name.lower()
#     return 'heading' in style_name or '标题' in style_name

# def get_heading_level(paragraph):
#     match = re.search(r'\d+', paragraph.style.name)
#     return int(match.group()) if match else 1

# def parse_paragraph(para, doc):
#     """解析段落为内容元素"""
#     element = {'type': 'text', 'content': '', 'images': []}
#     text_parts = []
    
#     for run in para.runs:
#         text_parts.append(run.text)
#         # 提取图片
#         if run._element.xpath('.//*[local-name()="blip"]'):
#             rel_id = run._element.xpath('.//*[local-name()="blip"]/@*[local-name()="embed"]')[0]
#             image_part = doc.part.related_parts[rel_id]
#             img_stream = io.BytesIO(image_part._blob)
#             element['images'].append(img_stream)
    
#     element['content'] = ''.join(text_parts).strip()
    
#     # 优先处理图片元素
#     if element['images']:
#         return {
#             'type': 'image',
#             'stream': element['images'][0],
#             'height': Inches(3)  # 固定图片高度
#         }
    
#     return element if element['content'] else None

# def parse_table(table):
#     """解析表格为内容元素"""
#     return {
#         'type': 'table',
#         'rows': len(table.rows),
#         'cols': len(table.columns),
#         'data': [[cell.text for cell in row.cells] for row in table.rows]
#     }


from docx import Document
from docx.oxml.shared import qn
from docx.oxml.ns import nsmap
import math
from docx import Document
from docx.oxml import parse_xml
import os
from docx import Document
import zipfile
import os
from docx.oxml.ns import nsmap
from lxml import etree

def delete_text_ranges_function(doc_path, output_path, paragraph_index, start_pos, end_pos):
    """处理文本删除范围"""
    doc = Document(doc_path)
    text_ranges = [(paragraph_index, start_pos, end_pos)]
    if not text_ranges:
        return
    
    # 建立段落索引映射
    para_dict = {i:p for i,p in enumerate(doc.paragraphs)}
    
    for para_idx, start, end in text_ranges:
        if para_idx not in para_dict:
            continue
            
        para = para_dict[para_idx]
        original_text = para.text
        
        # 处理非法位置
        start = max(0, min(start, len(original_text)))
        end = max(start, min(end, len(original_text)))
        
        # 生成新文本
        new_text = original_text[:start] + original_text[end:]
        para.text = new_text
        
        # 同步修改run对象
        sync_runs(para, original_text, new_text)
    doc.save(output_path)

def sync_runs(para, original, modified):
    """
    保持格式同步修改run对象
    算法逻辑：
    1. 计算修改前后的字符位置映射
    2. 重新分配run的文本内容
    """
    # 创建位置映射表
    diff = len(original) - len(modified)
    mapping = [i if i < len(modified) else math.inf 
              for i in range(len(original))]
    
    # 重新分配run内容
    new_runs = []
    current_pos = 0
    
    for run in para.runs:
        run_text = run.text
        valid_chars = []
        
        for c in run_text:
            if mapping[current_pos] != math.inf:
                valid_chars.append(c)
            current_pos += 1
            
        if valid_chars:
            new_run = run._element
            new_run.text = ''.join(valid_chars)
            new_runs.append(new_run)
    
    # 清空原有runs后添加新runs
    para._element.clear_content()
    for r in new_runs:
        para._element.append(r)


def delete_table_function(doc_path, output_path, table_index):
    """删除指定序号的表格"""
    doc = Document(doc_path)
    # if not table_indices:
    #     return
    
    # 转换序号并倒序处理
    # indices = [i-1 for i in table_indices if i > 0]
    # for idx in sorted(indices, reverse=True):
    #     if idx < len(doc.tables):
    table_index  = table_index - 1
    if table_index <0:
        table_index =0
    _delete_table(doc.tables[table_index])
    doc.save(output_path)


def _delete_table(table):
    """删除表格"""
    table._element.getparent().remove(table._element)

def delete_image_function(doc_path, output_path, image_index):
    """
    delete_image(doc_path='input.docx', output_path='output1.docx', image_index=1)
    """
    # 加载文档
    doc = Document(doc_path)
    
    # 获取文档的所有图片
    all_media = []
    rels = doc.part.rels
    image_count = 0

    for paragraph in doc.paragraphs:
        # 遍历段落中的所有运行(run)
        for run in paragraph.runs:
            # 检查运行中是否包含图片
            if run._element.xpath('.//pic:pic'):
                image_count += 1
                # 如果是第二张图片，删除它
                if image_count == image_index:
                    # 删除包含图片的运行
                    paragraph._p.remove(run._element)
    
    # 收集所有图片关系
    for rel in rels.values():
        if "image" in rel.reltype:
            all_media.append(('image', rel._target))
    
    # 验证索引有效性
    if image_index < 1 or image_index > len(all_media):
        raise ValueError(f"实际检测到图片数：{len(all_media)}，请求索引：{image_index}")
    
    # 获取要删除的图片关系
    _, target = all_media[image_index-1]
    
    # 删除图片关系
    for rel in list(rels.values()):
        if rel._target == target:
            doc.part.drop_rel(rel.rId)
    
    # 保存文档
    temp_path = output_path.replace(".docx", "_temp.docx")
    doc.save(temp_path)
    
    # 验证文档完整性
    with zipfile.ZipFile(temp_path, 'r') as z:
        if 'word/document.xml' not in z.namelist():
            raise IOError("生成文档结构损坏")
    
    os.replace(temp_path, output_path)
    return True

# delete_image(doc_path='input.docx', output_path='output1.docx', image_index=1)


import ast
import base64
import logging
import math
import re
import xml.etree.ElementTree as ET
from io import BytesIO
from typing import Dict, List

import backoff
import numpy as np
from PIL import Image
from requests.exceptions import SSLError
import openai
from openai import OpenAI
import random
# from google.api_core.exceptions import (
#     BadRequest,
#     InternalServerError,
#     InvalidArgument,
#     ResourceExhausted,
# )

# from mm_agents.accessibility_tree_wrap.heuristic_retrieve import (
#     filter_nodes,
# )
# from mm_agents.prompts import (
#     AGENTNET_SYSTEM_PROMPT,
# )
import json

# logger = logging.getLogger("desktopenv.agent")

# FINISH_WORD = "FINISH"
# WAIT_WORD = "WAIT"
# ENV_FAIL_WORD = "error_env"
# CALL_USER = "call_user"

# pure_text_settings = ["a11y_tree"]

# attributes_ns_ubuntu = "https://accessibility.windows.example.org/ns/attributes"
# attributes_ns_windows = "https://accessibility.windows.example.org/ns/attributes"
# state_ns_ubuntu = "https://accessibility.ubuntu.example.org/ns/state"
# state_ns_windows = "https://accessibility.windows.example.org/ns/state"
# component_ns_ubuntu = "https://accessibility.ubuntu.example.org/ns/component"
# component_ns_windows = "https://accessibility.windows.example.org/ns/component"
# value_ns_ubuntu = "https://accessibility.ubuntu.example.org/ns/value"
# value_ns_windows = "https://accessibility.windows.example.org/ns/value"
# class_ns_windows = "https://accessibility.windows.example.org/ns/class"
# # More namespaces defined in OSWorld, please check desktop_env/server/main.py

# # 定义一个函数来解析每个 action
# def parse_action(action_str):
#     try:
#         # 解析字符串为 AST 节点
#         node = ast.parse(action_str, mode='eval')

#         # 确保节点是一个表达式
#         if not isinstance(node, ast.Expression):
#             raise ValueError("Not an expression")

#         # 获取表达式的主体
#         call = node.body

#         # 确保主体是一个函数调用
#         if not isinstance(call, ast.Call):
#             raise ValueError("Not a function call")

#         # 获取函数名
#         if isinstance(call.func, ast.Name):
#             func_name = call.func.id
#         elif isinstance(call.func, ast.Attribute):
#             func_name = call.func.attr
#         else:
#             func_name = None

#         # 获取关键字参数
#         kwargs = {}
#         for kw in call.keywords:
#             key = kw.arg
#             # 处理不同类型的值，这里假设都是常量
#             if isinstance(kw.value, ast.Constant):
#                 value = kw.value.value
#             elif isinstance(kw.value, ast.Str):  # 兼容旧版本 Python
#                 value = kw.value.s
#             else:
#                 value = None
#             kwargs[key] = value

#         return {
#             'function': func_name,
#             'args': kwargs
#         }

#     except Exception as e:
#         print(f"Failed to parse action '{action_str}': {e}")
#         return None
    
def escape_single_quotes(text):
    # 匹配未转义的单引号（不匹配 \\'）
    pattern = r"(?<!\\)'"
    return re.sub(pattern, r"\\'", text)

# def text_to_json(text):
#     text = text.strip()
#     text = text.split("\n")
#     thought, action = "", ""
#     mode = None
#     for line in text:
#         if line.startswith("Thought:"):
#             mode = "thought"
#         elif line.startswith("Action:"):
#             mode = "action"
#         if line.strip() == "":
#             break
        
#         if mode == "thought":
#             thought+= line
#         elif mode == "action":
#             action += line
#     logger.info(f"Thought: {thought}")
#     logger.info(f"Action: {action}")
#     return thought.replace("Thought:", "").strip(), json.loads(action.replace("Action:", "").strip())

# def parse_action_qwen2vl_agentnet(text, factor, image_height, image_width):
#     logger.info(f"Parsing action: {text}")
#     thought_str, actions = text_to_json(text)
#     outputs  =[]
#     if type(actions) is list:
#         for action in actions:
#             action_type = action["action"]
#             action_inputs = action
#             outputs.append({
#                 "reflection": "",
#                 "thought": thought_str,
#                 "action_type": action_type,
#                 "action_inputs": action_inputs,
#                 "text": text
#             })
#     elif type(actions) is dict:
#         outputs.append({
#             "reflection": "",
#             "thought": thought_str,
#             "action_type": actions["action"],
#             "action_inputs": actions,
#             "text": text
#         })
#     else:
#         raise ValueError(f"Unknown action format: {actions}")
#     return outputs

# def parsing_response_to_pyautogui_code(responses, image_height: int, image_width:int, input_swap:bool=True) -> str:
import re
def escape_single_quotes(text):
    # 匹配未转义的单引号（不匹配 \\'）
    pattern = r"(?<!\\)'"
    return re.sub(pattern, r"\\'", text)


def parsing_response_to_pyautogui_code(responses, input_swap:bool=True) -> str:
    '''
    将M模型的输出解析为OSWorld中的action，生成pyautogui代码字符串
    参数:
        response: 包含模型输出的字典，结构类似于：
        {
            "action_type": "hotkey",
            "action_inputs": {
                "hotkey": "command v",
                "start_box": None,
                "end_box": None
            }
        }
    返回:
        生成的pyautogui代码字符串
    '''

    pyautogui_code = f"import pyautogui\nimport time\n"
    if isinstance(responses, dict):
        responses = [responses]
    for response_id, response in enumerate(responses):
        if "observation" in response:
            observation = response["observation"]
        else:
            observation = ""

        if "thought" in response:
            thought = response["thought"]
        else:
            thought = ""
        
        # if response_id == 0:
        #     pyautogui_code += f"'''\nObservation:\n{observation}\n\nThought:\n{thought}\n'''\n"
        # else:
        #     pyautogui_code += f"\ntime.sleep(3)\n"

        action_dict = response
        action_type = action_dict.get("action_type")
        action_inputs = action_dict.get("action_inputs", {})
        
        if action_type == "HOT_KEY":
            # Parsing hotkey action
            hotkey = action_inputs.get("value", "").strip()
            if hotkey:
                # Handle other hotkeys
                keys = hotkey.split(" ")  # Split the keys by space
                pyautogui_code += f"\ntime.sleep(0.5)\n"
                pyautogui_code += f"\npyautogui.hotkey({', '.join([repr(k) for k in keys])})"
            else:
                pyautogui_code += f"\n# Unrecognized hot key action parsing: {action_type}"
                logger.warning(f"Unrecognized hot key action parsing: {response}")
        elif action_type == "ENTER":
            pyautogui_code += f"\npyautogui.press('enter')"
        elif action_type == "INPUT":
            # Parsing typing action using clipboard
            content = action_inputs.get("value", "")
            content = escape_single_quotes(content)
            if content:
                if input_swap:
                    pyautogui_code += f"\nimport pyperclip"
                    pyautogui_code += f"\npyperclip.copy('{content.strip()}')"
                    pyautogui_code += f"\ntime.sleep(0.5)\n"
                    pyautogui_code += f"\npyautogui.hotkey('command', 'v')"
                    pyautogui_code += f"\ntime.sleep(0.5)\n"
                    if content.endswith("\n") or content.endswith("\\n"):
                        pyautogui_code += f"\npyautogui.press('enter')"
                else:
                    pyautogui_code += f"\npyautogui.write('{content.strip()}', interval=0.1)"
                    pyautogui_code += f"\ntime.sleep(0.5)\n"
                    if content.endswith("\n") or content.endswith("\\n"):
                        pyautogui_code += f"\npyautogui.press('enter')"

        
        elif action_type in ["DRAG", "SELECT"]:
            # Parsing drag or select action based on start and end_boxes
            points = action_inputs.get("position")
            assert len(points) == 2 and len(points[0]) == 2 and len(points[1]) == 2, "Drag or select action should have 2 points"
            start_box = points[0]
            end_box = points[1]
            
            if start_box and end_box:
                x1, y1 = start_box # [x, y] float number between 0 and 1
                # sx = round(float(x1) * image_width, 3)
                # sy = round(float(y1) * image_height, 3)
                sx = round(float(x1), 3)
                sy = round(float(y1), 3)
                x2, y2 = end_box # [x, y] float number between 0 and 1
                # ex = round(float(x2) * image_width, 3)
                # ey = round(float(y2) * image_height, 3)
                ex = round(float(x2), 3)
                ey = round(float(y2), 3)
                pyautogui_code += (
                    f"\npyautogui.moveTo({sx}, {sy})\n"
                    f"\npyautogui.dragTo({ex}, {ey}, duration=1.0)\n"
                )
            else:
                logger.warning(f"Parse failed! Drag or select action should have 2 points: {response}")
        elif action_type == "SCROLL":
            # Parsing scroll action
            start_box = action_inputs.get("position")
            if start_box:
                # x1, y1, x2, y2 = eval(start_box)  # Assuming box is in [x1, y1, x2, y2]
                # x = round(float((x1 + x2) / 2) * image_width, 3)
                # y = round(float((y1 + y2) / 2) * image_height, 3)
                x, y = start_box
                # x = round(float(x) * image_width, 3)
                # y = round(float(y) * image_height, 3)
                x = round(float(x), 3)
                y = round(float(y), 3)                
                # # 先点对应区域，再滚动
                pyautogui_code += f"\npyautogui.click({x}, {y}, button='left')"
            else:
                x = None
                y = None
            direction = action_inputs.get("direction", "")
            
            if x == None:
                if "up" in direction.lower():
                    pyautogui_code += f"\npyautogui.scroll(5)"
                elif "down" in direction.lower():
                    pyautogui_code += f"\npyautogui.scroll(-5)"
            else:
                if "up" in direction.lower():
                    pyautogui_code += f"\npyautogui.scroll(5, x={x}, y={y})"
                elif "down" in direction.lower():
                    pyautogui_code += f"\npyautogui.scroll(-5, x={x}, y={y})"

        elif action_type in ["CLICK", "LEFT_CLICK_DOUBLE", "LEFT_CLICK_SINGLE", "RIGHT_CLICK_SINGLE", "HOVER"]:
            # Parsing mouse click actions
            start_box = action_inputs.get("position")
            # start_box = str(start_box)
            if start_box:
                # start_box = eval(start_box)
                if len(start_box) == 2:
                    x1, y1 = start_box
                    x2 = x1
                    y2 = y1
                else:
                    raise ValueError(f"Unknown start_box format: {start_box}")
                # x = round(float((x1 + x2) / 2) * image_width, 3)
                # y = round(float((y1 + y2) / 2) * image_height, 3)

                x = round(float((x1 + x2)/2), 3)
                y = round(float((y1 + y2)/2), 3)

                if action_type == "LEFT_CLICK_DOUBLE" or action_type == "CLICK":
                    pyautogui_code += f"\npyautogui.click({x}, {y}, button='left')"
                elif action_type == "LEFT_CLICK_DOUBLE":
                    pyautogui_code += f"\npyautogui.doubleClick({x}, {y}, button='left')"
                elif action_type == "RIGHT_CLICK_SINGLE":
                    pyautogui_code += f"\npyautogui.click({x}, {y}, button='right')"
                elif action_type == "HOVER":
                    pyautogui_code += f"\npyautogui.moveTo({x}, {y})"
        
        elif action_type in ["FINISH"]:
            pyautogui_code = f"DONE"
        
        else:
            pyautogui_code += f"\n# Unrecognized action type: {action_type}"
            logger.error(f"Unrecognized action type: {response}")

    return pyautogui_code

# def pil_to_base64(image):
#     buffer = BytesIO()
#     image.save(buffer, format="PNG")  # 你可以改成 "JPEG" 等格式
#     return base64.b64encode(buffer.getvalue()).decode("utf-8")

# def linearize_accessibility_tree(accessibility_tree, platform="ubuntu"):

#     if platform == "ubuntu":
#         _attributes_ns = attributes_ns_ubuntu
#         _state_ns = state_ns_ubuntu
#         _component_ns = component_ns_ubuntu
#         _value_ns = value_ns_ubuntu
#     elif platform == "windows":
#         _attributes_ns = attributes_ns_windows
#         _state_ns = state_ns_windows
#         _component_ns = component_ns_windows
#         _value_ns = value_ns_windows
#     else:
#         raise ValueError("Invalid platform, must be 'ubuntu' or 'windows'")

#     filtered_nodes = filter_nodes(ET.fromstring(accessibility_tree), platform)
#     linearized_accessibility_tree = [
#         "tag\tname\ttext\tclass\tdescription\tposition (top-left x&y)\tsize (w&h)"
#     ]

#     # Linearize the accessibility tree nodes into a table format
#     for node in filtered_nodes:
#         if node.text:
#             text = (
#                 node.text
#                 if '"' not in node.text
#                 else '"{:}"'.format(node.text.replace('"', '""'))
#             )

#         elif node.get("{{{:}}}class".format(class_ns_windows), "").endswith(
#             "EditWrapper"
#         ) and node.get("{{{:}}}value".format(_value_ns)):
#             node_text = node.get("{{{:}}}value".format(_value_ns), "")
#             text = (
#                 node_text
#                 if '"' not in node_text
#                 else '"{:}"'.format(node_text.replace('"', '""'))
#             )
#         else:
#             text = '""'

#         linearized_accessibility_tree.append(
#             "{:}\t{:}\t{:}\t{:}\t{:}\t{:}\t{:}".format(
#                 node.tag,
#                 node.get("name", ""),
#                 text,
#                 (
#                     node.get("{{{:}}}class".format(_attributes_ns), "")
#                     if platform == "ubuntu"
#                     else node.get("{{{:}}}class".format(class_ns_windows), "")
#                 ),
#                 node.get("{{{:}}}description".format(_attributes_ns), ""),
#                 node.get("{{{:}}}screencoord".format(_component_ns), ""),
#                 node.get("{{{:}}}size".format(_component_ns), ""),
#             )
#         )

#     return "\n".join(linearized_accessibility_tree)

# def trim_accessibility_tree(linearized_accessibility_tree, max_tokens):
#     # enc = tiktoken.encoding_for_model("gpt-4")
#     # tokens = enc.encode(linearized_accessibility_tree)
#     # if len(tokens) > max_tokens:
#     #     linearized_accessibility_tree = enc.decode(tokens[:max_tokens])
#     #     linearized_accessibility_tree += "[...]\n"
#     return linearized_accessibility_tree

# class AgentNetAgent:
#     def __init__(
#         self,
#         platform="ubuntu",
#         max_tokens=1000,
#         top_p=0.3,
#         top_k=1.0,
#         temperature=0.0,
#         action_space="pyautogui",
#         observation_type="screenshot_a11y_tree",
#         # observation_type can be in ["screenshot", "a11y_tree", "screenshot_a11y_tree", "som"]
#         max_trajectory_length=50,
#         a11y_tree_max_tokens=10000,
#         runtime_conf: dict = {
#             "infer_mode": "qwen2vl_user",
#             "prompt_style": "qwen2vl_user",
#             "input_swap": True,
#             "language": "Chinese",
#             "max_steps": 50,
#             "history_n": 3,
#             "screen_height": 1080,
#             "screen_width": 1920
#         }
#     ):
#         self.platform = platform
#         self.max_tokens = max_tokens
#         self.top_p = top_p
#         self.top_k = top_k
#         self.temperature = temperature
#         self.action_space = action_space
#         self.observation_type = observation_type
#         self.max_trajectory_length = max_trajectory_length
#         self.a11y_tree_max_tokens = a11y_tree_max_tokens
#         self.runtime_conf = runtime_conf
#         self.vlm = OpenAI(
#             base_url="http://127.0.0.1:50004/v1",
#             api_key="empty",
#         ) # should replace with your UI-TARS server api
#         self.model_name = "agentnet"
#         self.infer_mode = self.runtime_conf["infer_mode"]
#         self.prompt_style = self.runtime_conf["prompt_style"]
#         self.input_swap = self.runtime_conf["input_swap"]
#         self.language = self.runtime_conf["language"]
#         self.max_steps = self.runtime_conf["max_steps"]

#         self.thoughts = []
#         self.actions = []
#         self.observations = []
#         self.history_images = []
#         self.history_responses = []
#         self.n_sampling = 1
#         self.system_prompt = AGENTNET_SYSTEM_PROMPT
#         self.customize_action_parser = parse_action_qwen2vl_agentnet
#         self.action_parse_res_factor = 1
#         if "history_n" in self.runtime_conf:
#             self.history_n = self.runtime_conf["history_n"]
#         else:
#             self.history_n = 5

#     def predict(
#         self, instruction: str, obs: Dict, last_action_after_obs: Dict = None
#     ) -> List:
#         """
#         Predict the next action(s) based on the current observation.
#         """

#         # Append trajectory
#         # print(len(self.observations), len(self.actions), len(self.actions))
#         assert len(self.observations) == len(self.actions) and len(self.actions) == len(
#             self.thoughts
#         ), "The number of observations and actions should be the same."

#         if len(self.observations) > self.max_trajectory_length:
#             if self.max_trajectory_length == 0:
#                 _observations = []
#                 _actions = []
#                 _thoughts = []
#             else:
#                 _observations = self.observations[-self.max_trajectory_length :]
#                 _actions = self.actions[-self.max_trajectory_length :]
#                 _thoughts = self.thoughts[-self.max_trajectory_length :]
#         else:
#             _observations = self.observations
#             _actions = self.actions
#             _thoughts = self.thoughts

#         for previous_obs, previous_action, previous_thought in zip(
#             _observations, _actions, _thoughts
#         ):
#             # {{{1
#             # print(previous_obs.keys())
#             if self.observation_type == "screenshot_a11y_tree":
#                 _screenshot = previous_obs["screenshot"]
#                 _linearized_accessibility_tree = previous_obs["accessibility_tree"]
#             elif self.observation_type == "screenshot":
#                 _screenshot = previous_obs["screenshot"]
#                 _linearized_accessibility_tree = None
#             else:
#                 raise ValueError(
#                     "Invalid observation_type type: " + self.observation_type
#                 )  # 1}}}

#         if last_action_after_obs is not None and self.infer_mode == "double_image":
#             self.history_images.append(last_action_after_obs["screenshot"])

#         self.history_images.append(obs["screenshot"])

#         if self.observation_type in ["screenshot", "screenshot_a11y_tree"]:
#             base64_image = obs["screenshot"]
#             try:
#                 linearized_accessibility_tree = (
#                     linearize_accessibility_tree(
#                         accessibility_tree=obs["accessibility_tree"],
#                         platform=self.platform,
#                     )
#                     if self.observation_type == "screenshot_a11y_tree"
#                     else None
#                 )
#             except:
#                 linearized_accessibility_tree = None
#             # logger.debug("LINEAR AT: %s", linearized_accessibility_tree)

#             if linearized_accessibility_tree:
#                 linearized_accessibility_tree = trim_accessibility_tree(
#                     linearized_accessibility_tree, self.a11y_tree_max_tokens
#                 )

#             if self.observation_type == "screenshot_a11y_tree":
#                 self.observations.append(
#                     {
#                         "screenshot": base64_image,
#                         "accessibility_tree": linearized_accessibility_tree,
#                     }
#                 )
#             else:
#                 self.observations.append(
#                     {"screenshot": base64_image, "accessibility_tree": None}
#                 )

#         else:
#             raise ValueError(
#                 "Invalid observation_type type: " + self.observation_type
#             )  # 1}}}
        
#         user_prompt = [
#             {
#                 "type": "text",
#                 "text": self.system_prompt
#             }, {
#                 "type": "text",
#                 "text": "Task: " + instruction
#             }
#         ]
#         if len(self.history_images) > self.history_n:
#             self.history_images = self.history_images[-self.history_n:]

#         max_pixels = 2000 * 28 * 28
#         min_pixels = 100 * 28 * 28
#         messages, images = [], []
#         if isinstance(self.history_images, bytes):
#             self.history_images = [self.history_images]
#         elif isinstance(self.history_images, np.ndarray):
#             self.history_images = list(self.history_images)
#         elif isinstance(self.history_images, list):
#             pass
#         else:
#             raise TypeError(f"Unidentified images type: {type(self.history_images)}")
#         max_image_nums_under_32k = int(32768*0.75/max_pixels*28*28)
#         if len(self.history_images) > max_image_nums_under_32k:
#             num_of_images = min(5, len(self.history_images))
#             max_pixels = int(32768*0.75) // num_of_images

#         for turn, image in enumerate(self.history_images):
#             if len(images) >= 5:
#                 break
#             try:
#                 image = Image.open(BytesIO(image))
#             except Exception as e:
#                 raise RuntimeError(f"Error opening image: {e}")

#             if image.width * image.height > max_pixels:
#                 """
#                 如果图片超过/低于像素限制，则计算一个缩放因子resize_factor，使图片的像素数缩小到等于或小于max_pixels。这个缩放因子是通过开平方根计算的，确保纵横比保持不变,这样原始的相对坐标可以不经转换直接复用
#                 """
#                 resize_factor = math.sqrt(max_pixels / (image.width * image.height))
#                 width, height = int(image.width * resize_factor), int(image.height * resize_factor)
#                 image = image.resize((width, height))
#             if image.width * image.height < min_pixels:
#                 resize_factor = math.sqrt(min_pixels / (image.width * image.height))
#                 width, height = math.ceil(image.width * resize_factor), math.ceil(image.height * resize_factor)
#                 image = image.resize((width, height))

#             if image.mode != "RGB":
#                 image = image.convert("RGB")

#             images.append(image)

#         messages = [
#             {
#                 "role": "system",
#                 "content": [{"type": "text", "text": "You are a helpful assistant."}]
#             },
#             {
#                 "role": "user",
#                 "content": user_prompt
#             }
#         ]
#         logger.info(f"history responses: {self.history_responses}; history images: {len(self.history_images)}")
#         image_num = 0
#         if len(self.history_responses) > 0:
#             for history_idx, history_response in enumerate(self.history_responses):
#                 # send at most history_n images to the model
#                 if history_idx + self.history_n > len(self.history_responses):

#                     cur_image = images[image_num]
#                     encoded_string = pil_to_base64(cur_image)
#                     # messages.append({
#                     #     "role": "user",
#                     #     "content": [{"type": "image_url", "image_url": {"url": f"data:image/png;base64,{encoded_string}"}}]
#                     # })
#                     messages[-1]["content"].append({
#                         "type": "image_url",
#                         "image_url": {"url": f"data:image/png;base64,{encoded_string}"}
#                     })
#                     image_num += 1
                    
#                     if type(history_response) == str:
#                         messages[-1]["content"].append({"type": "text", "text": history_response})
#                     else:
#                         messages[-1]["content"].append(history_response)

#             cur_image = images[image_num]
#             encoded_string = pil_to_base64(cur_image)
#             # messages.append({
#             #     "role": "user",
#             #     "content": [{"type": "image_url", "image_url": {"url": f"data:image/png;base64,{encoded_string}"}}]
#             # })
#             messages[-1]["content"].append({
#                 "type": "image_url",
#                 "image_url": {"url": f"data:image/png;base64,{encoded_string}"}
#             })
#             image_num += 1
        
#         else:
#             cur_image = images[image_num]
#             encoded_string = pil_to_base64(cur_image)
#             # messages.append({
#             #     "role": "user",
#             #     "content": [{"type": "image_url", "image_url": {"url": f"data:image/png;base64,{encoded_string}"}}]
#             # })
#             messages[-1]["content"].append({
#                 "type": "image_url",
#                 "image_url": {"url": f"data:image/png;base64,{encoded_string}"}
#             })
#             image_num += 1

#         try_times = 3
#         while True:
#             if try_times <= 0:
#                 logger.error(f"Reach max retry times to fetch response from client, as error flag.")
#                 return "client error", ["DONE"]
#             response = None
#             try:
#                 # logger.info("Call openai with prompt:")
#                 # import json
#                 # with open("debug.json", "w") as f:
#                 #     message_holder = []
#                 #     for message in messages:
#                 #         temp = {
#                 #             "role": message["role"],
#                 #             "content": message["content"] if type(message["content"]) == str else []
#                 #         }
#                 #         if type(message["content"]) == list:        
#                 #             for msg in message["content"]:
#                 #                 try:
#                 #                     if type(msg) == str:
#                 #                         temp["content"].append(msg)
#                 #                     else:
#                 #                         if msg["type"] == "image_url":
#                 #                             temp["content"].append({
#                 #                                 "type": "image_url",
#                 #                                 "image_url": "<image>"
#                 #                             })
#                 #                         else:
#                 #                             temp["content"].append({
#                 #                                 "type": msg["type"],
#                 #                                 "text": msg["text"]
#                 #                             })
#                 #                 except Exception as e:
#                 #                     logger.error(f"Error when parsing message: {e}")
#                 #                     # print(msg)
#                 #                     # print(msg)
#                 #                     # exit()
#                 #         message_holder.append(temp)
                    
#                 #     logger.info(f"{json.dumps(message_holder, indent=4, ensure_ascii=False)}")
#                 #     json.dump(message_holder, f, indent=4, ensure_ascii=False)
                    
#                 response = self.vlm.chat.completions.create(
#                     model=self.model_name,
#                     messages=messages,
#                     frequency_penalty=1,
#                     max_tokens=self.max_tokens,
#                     temperature=self.temperature,
#                     top_p=self.top_p,
#                     n=self.n_sampling
#                 )
#                 # logger.info(f"response: \n{response.choices}")
#                 random_idx = random.randint(0, self.n_sampling - 1)
#                 prediction = response.choices[random_idx].message.content.strip()
#                 logger.info(f"agentnet response: \n{prediction}")
#                 # prediction = response[0]["prediction"].strip()
#                 parsed_responses = self.customize_action_parser(
#                     prediction,
#                     self.action_parse_res_factor,
#                     self.runtime_conf["screen_height"],
#                     self.runtime_conf["screen_width"]
#                 )
#                 break
#             except Exception as e:
#                 print(f"Error when fetching response from client, with response: {response}")
#                 import traceback
#                 traceback.print_exc()
#                 # exit()
#                 prediction = None
#                 try_times -= 1
                
#         if prediction is None:
#             logger.error("Prediction is None, return DONE")
#             return "client error", ["DONE"]

        
#         self.history_responses.append(prediction)
#         self.thoughts.append(prediction)

#         try:
#             parsed_responses = self.customize_action_parser(
#                 prediction,
#                 self.action_parse_res_factor,
#                 self.runtime_conf["screen_height"],
#                 self.runtime_conf["screen_width"]
#             )
#         except Exception as e:
#             print(f"Parsing action error: {prediction}, with error:\n{e}")
#             return f"Parsing action error: {prediction}, with error:\n{e}", ["DONE"]

#         actions = []
#         for parsed_response in parsed_responses:
#             if "action_type" in parsed_response:

#                 if parsed_response["action_type"] == FINISH_WORD:
#                     self.actions.append(actions)

#                     return prediction, ["DONE"]
                
#                 elif parsed_response["action_type"] == WAIT_WORD:
#                     self.actions.append(actions)
#                     return prediction, ["WAIT"]
                
#                 elif parsed_response["action_type"] == ENV_FAIL_WORD:
#                     self.actions.append(actions)
#                     return prediction, ["FAIL"]

#                 elif parsed_response["action_type"] == CALL_USER:
#                     self.actions.append(actions)
#                     return prediction, ["FAIL"]

            
#             pyautogui_code = parsing_response_to_pyautogui_code(
#                 parsed_response,
#                 self.runtime_conf["screen_height"],
#                 self.runtime_conf["screen_width"],
#                 self.input_swap
#             )
#             actions.append(pyautogui_code)

#         self.actions.append(actions)

#         if len(self.history_responses) >= self.max_trajectory_length:
#             # Default to FAIL if exceed max steps
#             actions = ["FAIL"]

#         return prediction, actions


#     @backoff.on_exception(
#         backoff.constant,
#         # here you should add more model exceptions as you want,
#         # but you are forbidden to add "Exception", that is, a common type of exception
#         # because we want to catch this kind of Exception in the outside to ensure each example won't exceed the time limit
#         (
#             # General exceptions
#             SSLError,
#             # OpenAI exceptions
#             openai.RateLimitError,
#             openai.BadRequestError,
#             openai.InternalServerError,
#             # Google exceptions
#             InvalidArgument,
#             ResourceExhausted,
#             InternalServerError,
#             BadRequest,
#             # Groq exceptions
#             # todo: check
#         ),
#         interval=30,
#         max_tries=10,
#     )
    
#     def reset(self, runtime_logger):
#         self.thoughts = []
#         self.actions = []
#         self.observations = []
#         self.history_images = []
#         self.history_responses = []


from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.oxml import OxmlElement
# from docx.oxml.shared import qn
from docx.text.run import Run
from pathlib import Path
from typing import Union, Tuple, List, Dict
import docx.parts.image
from docx.oxml import parse_xml
import mimetypes
from docx.oxml.ns import nsmap, qn


from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
import os

def insert_image_function(doc_path, output_path, image_path, insert_position=0):
    """
    用法：在第二段后插入图片，格式自动调整
    doc_path: 文档的目录
    output_path: 处理后文档的目录
    image_path: 图片的目录
    insert_position: 插入在哪一段后
    success = insert_image(
        doc_path='input.docx',
        output_path="safe_output.docx",
        image_path="logo.png",
        insert_position=2
    )
    if success:
        print("请手动验证生成文件")
    """
    try:
        # 验证图片文件完整性
        if not os.path.exists(image_path):
            raise FileNotFoundError(f"图片文件不存在：{image_path}")
        with Image.open(image_path) as img:
            img.verify()  # 校验图片完整性
            dpi = img.info.get('dpi', (96, 96))[0]
            dpi = 96 if dpi == 0 else dpi
            img_width_inch = img.width / dpi
            img_height_inch = img.height / dpi
        doc = Document(doc_path) if doc_path else Document()
        section = doc.sections[0]
        page_width = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
        page_height = section.page_height.inches - section.top_margin.inches - section.bottom_margin.inches
        
        # 自适应缩放
        ratio = min(page_width / img_width_inch, page_height / img_height_inch)
        if ratio < 1:
            img_width_inch *= ratio
            img_height_inch *= ratio
        
        # 初始化文档
        doc = Document(doc_path) if os.path.exists(doc_path) else Document()
        
        # 创建新段落（标准方法）
        doc.add_paragraph()  # 前导空段落
        target_para = doc.add_paragraph()
        
        # 插入图片（使用python-docx原生方法）
        run = target_para.add_run()

        paragraph_format = target_para.paragraph_format
        paragraph_format.space_before = Pt(0)  # 段前间距
        paragraph_format.space_after = Pt(0)   # 段后间距
        paragraph_format.line_spacing = 1.0    # 单倍行距
        run.add_picture(image_path, 
                   width=Inches(img_width_inch), 
                   height=Inches(img_height_inch))
        
        # 设置段落格式
        target_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 移动段落到指定位置
        if insert_position < len(doc.paragraphs):
            para_element = target_para._element
            parent = doc.paragraphs[insert_position]._element.getparent()
            parent.insert(insert_position, para_element)
        
        # 保存文档（强制XML规范）
        temp_path = output_path.replace(".docx", "_temp.docx")
        doc.save(temp_path)
        
        # 二次验证文件完整性
        if os.path.getsize(temp_path) < 1024:
            raise IOError("生成文件过小，可能已损坏")
        os.replace(temp_path, output_path)
        
        # print(f"文件已安全保存至：{output_path}")
        return True

    except Exception as e:
        # print(f"生成失败：{str(e)}")
        if os.path.exists(temp_path):
            os.remove(temp_path)
        return False

def set_table_borders(table):
    """强制设置表格边框（全边框）"""
    tbl = table._tbl
    tblBorders = OxmlElement('w:tblBorders')
    
    # 定义所有边框类型
    borders = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
    for border_type in borders:
        border = OxmlElement(f'w:{border_type}')
        border.set(qn('w:val'), 'single')    # 实线
        border.set(qn('w:sz'), '4')          # 线宽4单位
        border.set(qn('w:space'), '0')       # 无间隔
        border.set(qn('w:color'), '000000')  # 黑色
        tblBorders.append(border)
    
    # 应用边框到表格属性
    tblPr = tbl.tblPr
    tblPr.append(tblBorders)

def get_valid_table_style(doc):
    """获取当前文档可用的表格样式"""
    # 内置的通用样式列表（兼容多语言）
    fallback_styles = [
        'Table Grid',        # 英文版
        '表格网格',          # 中文版
        'Light Shading',
        '浅色底纹',
        'Normal Table',
        '普通表格'
    ]
    
    for style in fallback_styles:
        if style in doc.styles:
            return style
    # print('None')
    return None  # 无可用样式时返回None

def insert_table_function(doc_path, target_para, rows, cols, data=None):
    """
    doc_path: doc文档的目录
    target_para: 目标段落(段落索引)
    rows: 表格的行数
    cols: 表格的列数
    data: 表格数据(list形式)
    使用示例：
    data = [
        ["Header1", "Header2", "Header3", "Header4"],
        ["内容1", "内容2", "内容3", "内容4"],
        ["合并单元格示例", "", "合并", ""]
    ]
    
    table = insert_table_safely(
        doc_path='input.docx',
        target_para=0,  # 也可以直接传段落索引如 1
        rows=3,
        cols=4,
        data=data,
        # style='Light Shading'
    )
    """
    doc = Document(doc_path) if doc_path else Document()
    # 创建表格结构（省略样式参数）
    tbl = OxmlElement('w:tbl')
    
    # 创建表格必须的结构元素
    tblPr = OxmlElement('w:tblPr')
    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(cols):
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), '5000')  # 默认列宽5cm
        tblGrid.append(gridCol)
    tbl.append(tblPr)
    tbl.append(tblGrid)
    
    # 创建行和列
    for _ in range(rows):
        tr = OxmlElement('w:tr')
        for _ in range(cols):
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), '5000')
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            tc.append(tcPr)
            tr.append(tc)
        tbl.append(tr)
    
    # 插入到文档
    if isinstance(target_para, int):
        para = doc.paragraphs[target_para]
    else:
        para = target_para
    para._p.addnext(tbl)
    
    # 获取有效样式
    table = doc.tables[-1]
    valid_style = get_valid_table_style(doc)
    if valid_style:
        try:
            table.style = valid_style
        except KeyError:
            pass  # 样式应用失败时保持无样式
    # 填充数据
    if data:
        for row_idx in range(min(rows, len(data))):
            for col_idx in range(min(cols, len(data[row_idx]))):
                table.cell(row_idx, col_idx).text = str(data[row_idx][col_idx])
    table = set_table_borders(table)
    doc.save("table_inserted.docx")
    return table



# from docx import Document

def create_docx_function(save_path):
    doc = Document()
    doc.save(save_path)
